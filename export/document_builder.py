import io
import re
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal line (paragraph border) to the document."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_blockquote_paragraph(doc: Document, text: str) -> None:
    """Add a blockquote-styled paragraph with a left border."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Add left border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "4472C4")
    pBdr.append(left)
    pPr.append(pBdr)


def _parse_and_add_answer(doc: Document, answer_text: str) -> None:
    """
    Parse Claude's markdown-style answer and add formatted paragraphs to the document.
    Handles:
      - Lines starting with '>' → blockquote style
      - **bold** text → bold runs
      - Regular text → normal paragraphs
    """
    lines = answer_text.split("\n")
    buffer = []

    def flush_buffer():
        if buffer:
            combined = " ".join(buffer).strip()
            if combined:
                _add_formatted_paragraph(doc, combined)
            buffer.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_buffer()
            continue

        if stripped.startswith("> "):
            flush_buffer()
            _add_blockquote_paragraph(doc, stripped[2:])
        elif stripped.startswith(">"):
            flush_buffer()
            _add_blockquote_paragraph(doc, stripped[1:].strip())
        else:
            buffer.append(stripped)

    flush_buffer()


def _add_formatted_paragraph(doc: Document, text: str) -> None:
    """Add a paragraph with **bold** markdown rendered as bold runs."""
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def build_docx(
    session_title: str,
    uploaded_files: list[str],
    concepts: dict,
    qa_history: list[dict],
) -> bytes:
    """
    Build a Word document (.docx) with:
    1. Title page
    2. Key Concepts table
    3. Questions & Answers
    4. References

    Returns bytes suitable for st.download_button.
    """
    doc = Document()

    # --- Title Page ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(session_title)
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Co-Creation in Education — Decision Support Report")
    sub.runs[0].font.size = Pt(13)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")

    doc.add_paragraph()

    if uploaded_files:
        doc.add_heading("Literature Used", level=2)
        for fname in uploaded_files:
            doc.add_paragraph(f"• {fname}", style="List Bullet")

    doc.add_page_break()

    # --- Key Concepts ---
    all_concepts = []
    for fname, clist in concepts.items():
        all_concepts.extend(clist)

    if all_concepts:
        doc.add_heading("Key Concepts from the Literature", level=1)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Term"
        hdr[1].text = "Definition"
        hdr[2].text = "Source"

        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True

        for concept in all_concepts:
            row = table.add_row().cells
            row[0].text = concept.get("term", "")
            row[1].text = concept.get("definition", "")
            row[2].text = concept.get("source_filename", "")

        doc.add_paragraph()
        doc.add_page_break()

    # --- Questions & Answers ---
    doc.add_heading("Questions and Answers", level=1)

    if not qa_history:
        doc.add_paragraph("No questions have been asked in this session.")
    else:
        for i, qa in enumerate(qa_history, start=1):
            doc.add_heading(f"Question {i}", level=2)
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(qa["question"])
            q_run.bold = True
            q_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

            doc.add_paragraph()
            doc.add_heading("Answer", level=3)
            _parse_and_add_answer(doc, qa["answer"])

            doc.add_paragraph()
            ts_para = doc.add_paragraph()
            ts_para.add_run(f"Asked: {qa.get('timestamp', '')}").font.color.rgb = RGBColor(
                0x88, 0x88, 0x88
            )
            _add_horizontal_rule(doc)
            doc.add_paragraph()

    # --- References ---
    if uploaded_files:
        doc.add_page_break()
        doc.add_heading("References", level=1)
        for i, fname in enumerate(uploaded_files, start=1):
            doc.add_paragraph(f"{i}. {fname}", style="List Number")

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
